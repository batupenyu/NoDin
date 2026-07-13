from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class NotaDinasFiller:
    def __init__(self, template_path="nota_dinas_template_fillable.docx"):
        self.template_path = template_path
        self.context = {}

    @staticmethod
    def _set_run_font(run, name="Arial", size_pt=12):
        run.font.name = name
        run.font.size = Pt(size_pt)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(int(size_pt * 2)))

    @classmethod
    def _apply_font(cls, doc, name="Arial", size_pt=12):
        for para in doc.paragraphs:
            for run in para.runs:
                cls._set_run_font(run, name, size_pt)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            cls._set_run_font(run, name, size_pt)

    def fill_header(self, Kepada="Kepala Sekolah SMKN 1 Koba",
                    Dari="SMK Negeri 1 Koba",
                    Nomor="", Lampiran="-", Tanggal="", Perihal=""):
        self.context['KEPADA'] = Kepada
        self.context['DARI'] = Dari
        self.context['NOMOR_SURAT'] = Nomor
        self.context['LAMPIRAN'] = Lampiran
        self.context['TANGGAL_SURAT'] = Tanggal
        self.context['PERIHAL'] = Perihal

    def fill_kegiatan(self, kegiatan=""):
        self.context['NAMA_KEGIATAN'] = kegiatan

    def fill_dasar_pelaksanaan(self, nomor=""):
        self.context['DASAR_PELAKSANAAN'] = f"800.1.11.1/{nomor}/SMKN 1 Kb/Dindik/2026"

    def fill_maksud_tujuan(self, text=""):
        self.context['MAKSUD_TUJUAN'] = text

    def fill_peserta(self, peserta_list):
        for i, p in enumerate(peserta_list[:2]):
            self.context[f'PESERTA{i+1}_NAMA'] = p.get('nama', '')
            self.context[f'PESERTA{i+1}_NIP'] = p.get('nip', '')
            self.context[f'PESERTA{i+1}_JABATAN'] = p.get('jabatan', '')
        self.context['peserta_list'] = peserta_list

    def fill_pelaksanaan(self, tanggal="", jam="", tempat=""):
        self.context['TGL_PELAKSANAAN'] = tanggal
        self.context['JAM_PELAKSANAAN'] = jam
        self.context['TEMPAT_PELAKSANAAN'] = tempat

    def fill_kesimpulan(self, text=""):
        self.context['KESIMPULAN'] = text

    def fill_tembusan(self, entries):
        for i, e in enumerate(entries[:2]):
            self.context[f'tem_nama_{i+1}'] = e.get('nama', '')
            self.context[f'tem_nip_{i+1}'] = e.get('nip', '')
            self.context[f'tem_jabatan_{i+1}'] = e.get('jabatan', '')

    def save(self, output_path):
        doc = DocxTemplate(self.template_path)
        doc.render(self.context)
        peserta_list = self.context.get('peserta_list', [])
        if len(peserta_list) <= 1:
            self._remove_peserta2(doc)
        self._apply_font(doc, "Arial", 12)
        doc.save(output_path)

    @staticmethod
    def _remove_peserta2(doc):
        from docx.oxml.ns import qn

        t0 = doc.tables[0]
        row2 = t0.rows[2]

        cell3 = row2.cells[3]
        for idx in [6,5,4]:
            if idx < len(cell3.paragraphs):
                cell3.paragraphs[idx]._element.getparent().remove(cell3.paragraphs[idx]._element)

        cell4 = row2.cells[4]
        for idx in [6,5,4]:
            if idx < len(cell4.paragraphs):
                cell4.paragraphs[idx]._element.getparent().remove(cell4.paragraphs[idx]._element)

        cell5 = row2.cells[5]
        for idx in [6,5,4]:
            if idx < len(cell5.paragraphs):
                cell5.paragraphs[idx]._element.getparent().remove(cell5.paragraphs[idx]._element)

        t1 = doc.tables[1]
        row0 = t1.rows[0]

        cell2 = row0.cells[2]
        if 7 < len(cell2.paragraphs):
            cell2.paragraphs[7]._element.getparent().remove(cell2.paragraphs[7]._element)

        cell3 = row0.cells[3]
        for idx in [9,8,7]:
            if idx < len(cell3.paragraphs):
                cell3.paragraphs[idx]._element.getparent().remove(cell3.paragraphs[idx]._element)

        cell4 = row0.cells[4]
        for idx in [9,8,7]:
            if idx < len(cell4.paragraphs):
                cell4.paragraphs[idx]._element.getparent().remove(cell4.paragraphs[idx]._element)

        cell5 = row0.cells[5]
        for idx in [9,8,7]:
            if idx < len(cell5.paragraphs):
                cell5.paragraphs[idx]._element.getparent().remove(cell5.paragraphs[idx]._element)
