import re
from docx import Document
from docx.oxml.ns import qn

d = Document("nota_dinas.docx")

print("="*70)
print("ALL PARAGRAPH TEXT")
print("="*70)
for i,p in enumerate(d.paragraphs):
    txt = p.text
    if txt.strip():
        print(f"[P{i}] {txt}")

print()
print("="*70)
print("ALL TABLE CONTENTS")
print("="*70)
for ti,t in enumerate(d.tables):
    print(f"--- TABLE {ti} ({len(t.rows)}x{len(t.columns)}) ---")
    for ri,row in enumerate(t.rows):
        cells = [c.text.replace("\n"," | ") for c in row.cells]
        print(f"  R{ri}: " + " || ".join(cells))
