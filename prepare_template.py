import os
from docx import Document
from docx.oxml.ns import qn


def prepare_tagged_template(src_path, dst_path):
    doc = Document(src_path)

    def tag_para_after_colon(para, placeholder):
        full = para.text
        idx = full.find(':')
        if idx == -1:
            return
        prefix = full[:idx + 1]
        para.clear()
        if para.runs:
            para.runs[0].text = f"{prefix}\t{placeholder}"
        else:
            para.add_run(f"{prefix}\t{placeholder}")

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt.startswith("Kepada"):
            tag_para_after_colon(para, "{{KEPADA}}")
        elif txt.startswith("Dari"):
            tag_para_after_colon(para, "{{DARI}}")
        elif txt.startswith("Nomor"):
            tag_para_after_colon(para, "{{NOMOR_SURAT}}")
        elif txt.startswith("Lampiran"):
            tag_para_after_colon(para, "{{LAMPIRAN}}")
        elif txt.startswith("Tanggal") and "Perihal" not in txt:
            tag_para_after_colon(para, "{{TANGGAL_SURAT}}")
        elif txt.startswith("Perihal"):
            tag_para_after_colon(para, "{{PERIHAL}}")

    tables = doc.tables

    table1 = tables[1]
    tem_cell = table1.rows[0].cells[4]
    tem_paras = tem_cell.paragraphs
    if len(tem_paras) > 9:
        for para in tem_paras[1:]:
            for run in para.runs:
                if run.text.strip() == ":":
                    run.text = ""
        tem_paras[1].add_run("{{tem_nama_1}}")
        tem_paras[2].add_run("{{tem_nip_1}}")
        tem_paras[3].add_run("{{tem_jabatan_1}}")
        tem_paras[7].add_run("{{tem_nama_2}}")
        tem_paras[8].add_run("{{tem_nip_2}}")
        tem_paras[9].add_run("{{tem_jabatan_2}}")

    table0 = doc.tables[0]
    row2_t0 = table0.rows[2]

    cell3_t0 = row2_t0.cells[3]
    if len(cell3_t0.paragraphs) >= 7:
        cell3_t0.paragraphs[4].text = '{% if peserta_list|length > 1 %}Nama{% endif %}'
        cell3_t0.paragraphs[5].text = '{% if peserta_list|length > 1 %}NIP{% endif %}'
        cell3_t0.paragraphs[6].text = '{% if peserta_list|length > 1 %}Jabatan  {% endif %}'

    cell5_t0 = row2_t0.cells[5]
    if len(cell5_t0.paragraphs) >= 7:
        cell5_t0.paragraphs[4].text = '{% if peserta_list|length > 1 %}{{PESERTA2_NAMA}}{% endif %}'
        cell5_t0.paragraphs[5].text = '{% if peserta_list|length > 1 %}{{PESERTA2_NIP}}{% endif %}'
        cell5_t0.paragraphs[6].text = '{% if peserta_list|length > 1 %}{{PESERTA2_JABATAN}}{% endif %}'

    table1 = doc.tables[1]
    row0_t1 = table1.rows[0]

    cell2_t1 = row0_t1.cells[2]
    if len(cell2_t1.paragraphs) >= 8:
        cell2_t1.paragraphs[7].text = '{% if peserta_list|length > 1 %}2.{% endif %}'

    cell3_t1 = row0_t1.cells[3]
    if len(cell3_t1.paragraphs) >= 10:
        cell3_t1.paragraphs[7].text = '{% if peserta_list|length > 1 %}Nama{% endif %}'
        cell3_t1.paragraphs[8].text = '{% if peserta_list|length > 1 %}NIP{% endif %}'
        cell3_t1.paragraphs[9].text = '{% if peserta_list|length > 1 %}Jabatan  {% endif %}'

    cell5_t1 = row0_t1.cells[5]
    if len(cell5_t1.paragraphs) >= 10:
        cell5_t1.paragraphs[7].text = '{% if peserta_list|length > 1 %}{{PESERTA2_NAMA}}{% endif %}'
        cell5_t1.paragraphs[8].text = '{% if peserta_list|length > 1 %}{{PESERTA2_NIP}}{% endif %}'
        cell5_t1.paragraphs[9].text = '{% if peserta_list|length > 1 %}{{PESERTA2_JABATAN}}                      {% endif %}'

    doc.save(dst_path)


def set_run_font(run, name="Arial", size_pt=12):
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run.font.name = name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(size_pt * 2)))


def set_document_default_font(doc, name="Arial", size_pt=12):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    body = doc.element.body
    rPrDefault = body.find('.//' + qn('w:docDefaults') + '//' + qn('w:rPrDefault'))
    if rPrDefault is None:
        docDefaults = body.find(qn('w:docDefaults'))
        if docDefaults is None:
            docDefaults = OxmlElement('w:docDefaults')
            body.insert(0, docDefaults)
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(size_pt * 2)))


def apply_font_to_document(doc, name="Arial", size_pt=12):
    set_document_default_font(doc, name, size_pt)
    for para in doc.paragraphs:
        for run in para.runs:
            set_run_font(run, name, size_pt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, name, size_pt)


if __name__ == "__main__":
    src = os.path.join(os.path.dirname(__file__), "nota_dinas_TEMPLATE.docx")
    dst = os.path.join(os.path.dirname(__file__), "nota_dinas_template_fillable.docx")
    prepare_tagged_template(src, dst)
    apply_font_to_document(Document(dst), "Arial", 12)
    Document(dst).save(dst)
    print(f"Prepared template saved to {dst}")
